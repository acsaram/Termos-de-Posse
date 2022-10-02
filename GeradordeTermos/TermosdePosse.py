from docx import Document
from docx.shared import Pt
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

document = Document()

# Estilos do texto

styles = document.styles
s = styles.add_style("Sublinhado", WD_STYLE_TYPE.PARAGRAPH)
s.paragraph_format.space_after = Pt(0)

# Variaveis

numero_termo_de_posse = "1104"
ano_vigente = str(datetime.now().year)
mes_vigente = "primeiro dia do mês de setembro de 2022 (01/09/2022)"
nome_diretor = "Ítalo Estrela Porto"
nome_posse = "Andrey Rocha Spago"
estado_civil = "Solteiro"
data_posse = "01/09/2022"
cargo = "ASSESSOR PARLAMENTAR DE GABINETE II"
simbolo_cargo = "APG-2"
numero_portaria = "1.154"
data_portaria =  "08/09/2022"
documento_id = '5556022 2ª via'
cpf = "751.263.201-06"

#Titulo e seus estilos

titulo = document.add_paragraph('TERMO DE POSSE Nº {}/{} \n\n'.format(numero_termo_de_posse, ano_vigente))
titulo.alignment = 1
titulo.style = document.styles.add_style('Titulo', WD_STYLE_TYPE.PARAGRAPH)
font = titulo.style.font
font.size = Pt(16)

# Paragrafo

p = document.add_paragraph('Ao {}, nesta cidade de Goiânia, capital do Estado de Goiás, na sala destinada à Diretoria de Recursos Humanos '
                           'deste Poder Legislativo Municipal, onde achava-se presente o(a) diretor(a) {}, compareceu o(a) Sr.(a) '.format(mes_vigente, nome_diretor))
p.add_run('{}, {} '.format(nome_posse.upper(), estado_civil)).bold = True
p.add_run('residente e domiciliado (a) nesta Capital, a quem o(a) Senhor(a) Diretor(a) deferiu a posse, após ouvir o seu compromisso de bem e fielmente desempenhar, a partir do dia')
p.add_run(' {}, '.format(data_posse)).bold = True
p.add_run('o cargo em comissão de ')
p.add_run('{}, SÍMBOLO {}'.format(cargo, simbolo_cargo)).bold = True
p.add_run(' do Quadro Próprio deste Poder Legislativo, constante da Lei n° 10.801, de 15 de julho de 2022, para o qual foi nomeado(a) pela')
p.add_run('Portaria nº {} de {}'.format(numero_portaria, data_portaria)).bold = True
p.add_run('. Apresentou o(a) empossado(a) os documentos de identidade n° ')
p.add_run('{} '.format(documento_id)).bold = True
p.add_run('e CPF nº ')
p.add_run('{}.'.format(cpf)).bold = True
p.add_run(' Assinou declaração o de não exercício de outro cargo, emprego ou função pública, declarou que não é parente, até o terceiro grau, de qualquer membro deste Poder ou de servidores investidos em cargos'
                        'de direção ou assessoramento, em linha reta, colateral ou por afinidade e entregou declaração de atendimento ao artigo 20-A, da Lei Orgânica do Município '
                        'de Goiânia. Aceitando os termos da posse, prometeu cumprir com o máximo zelo e dedicação as mencionadas funções e declarou ter pleno conhecimento do '
                        'Estatuto dos Servidores Públicos da Câmara Municipal de Goiânia, Lei Complementar n° 354, de 15 de julho de 2022, entrando imediatamente no exercício do cargo.'
                        ' Para constar lavrou-se o presente termo, que lido e achado conforme, vai devidamente assinado pelo Senhor(a) Diretor(a) de Recursos Humanos e pelo empossado(a).')

p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

cidade = document.add_paragraph('\nGoiânia,         /               /{}'.format(ano_vigente))

sublinhado = document.add_paragraph('\n\n\n\n_____________________________________________', style = "Sublinhado")
empossado = document.add_paragraph('                         Empossado')


sublinhado2 = document.add_paragraph('\n\n\n\n_____________________________________________', style = "Sublinhado")
assinatura_diretor = document.add_paragraph('                           Diretor')



document.save('Termo de Posse nº {} - {}.docx'.format(numero_termo_de_posse, nome_posse))